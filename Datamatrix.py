#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Feb 18 23:40:23 2020

@author: eugeniy
"""
from pylibdmtx.pylibdmtx import encode
from PIL import Image
from docx import Document
from docx.shared import Mm

def marks_download (data=''):
    marks = [['123123323133dsfsdds3213ssdcfsdvsd12312','Обувь мужская','М10','Импорт'],
             ['12312332sd234sdfsdf3dsfsdfsdfsdfsdfsd','Обувь женская','М10','Импорт']]
    return marks

def page_generation (doc, mark):
    dtmx = get_image(mark[0])
    doc.add_picture(dtmx) 
    doc.add_paragraph(mark[1])
    doc.add_paragraph(mark[2])
    doc.add_paragraph(mark[3])
    doc.add_page_break()
        
def docx_file_genetation ():
    document = Document()
    document.add_heading('Document Title', 0)
    document.settings
    section = document.sections[0]
    section.page_height = Mm(58)
    section.page_width = Mm(60)
    section.left_margin = Mm(2)
    section.right_margin = Mm(2)
    section.top_margin = Mm(2)
    section.bottom_margin = Mm(2)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)
    marks=marks_download()
    for mark in marks:
        page_generation(document, mark)
    document.add_page_break()
    document.save('demo.docx')

def get_image(data ='Some text'): 
    encoded = encode(data)
    img = Image.frombytes('RGB', (encoded.width, encoded.height), encoded.pixels)
    filename = data+'.png'
    img.save(filename)
    return filename
    
docx_file_genetation()
    
